from llama_index.core import (
    Document,
    StorageContext,
    VectorStoreIndex,
    load_index_from_storage,
)
import os
from PyPDF2 import PdfReader
import docx
from llama_index.core.node_parser import SimpleNodeParser

# Function to manually load documents from a directory
def load_documents_from_directory(directory_path: str):
    documents = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if filename.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
                documents.append(Document(text=content, metadata={"filename": filename}))
        elif filename.endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
            documents.append(Document(text=text, metadata={"filename": filename}))
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(file_path)
            documents.append(Document(text=text, metadata={"filename": filename}))
    return documents

# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

# Function to extract text from DOCX
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

# Function to create and save a new index
def process_files_with_llama_index(directory_path: str, output_dir: str):
    """
    Processes files in the specified directory and saves the index in the output directory.
    """
    # Load the documents manually from the directory
    documents = load_documents_from_directory(directory_path)

    # Create a VectorStore index
    index = VectorStoreIndex.from_documents(documents)

    # Save the index in the storage context
    index.storage_context.persist(output_dir)


# Function to load and use an existing index
def load_index_and_query(storage_dir: str, query: str, retrieval_len: int):
    """
    Load the index from the storage directory and query it. 
    Returns a parsed response with relevant information.
    """
    # Load the index from the storage
    storage_context = StorageContext.from_defaults(persist_dir=storage_dir)
    index = load_index_from_storage(storage_context)
    
    # Convert index to retriever
    retriever = index.as_retriever()
    
    # Perform the retrieval
    results = retriever.retrieve(query)
    
    # Parse and format the results
    parsed_results = []
    text_result=[]
    for result in results[:retrieval_len]:
        node = result.node
        text_result.append(node.text)
        parsed_results.append({
            "text": node.text,  # The actual content of the node
            "metadata": node.metadata,  # Metadata like filename
            "score": result.score  # Relevance score
        })
    return parsed_results,text_result

# load_index_and_query("uploads/9657b257-f1fb-4479-ba25-683312e2e7ab/8cb5c2df-bbc7-49dc-a8af-3539a40963a1","total cost",13)